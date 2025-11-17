from flask import Flask, request, jsonify, render_template_string, send_file
from flask_cors import CORS
from openai import OpenAI
import os
import requests
import json
import re
import base64
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from PIL import Image
import io
import urllib.request

app = Flask(__name__)

# –ù–∞–ª–∞—à—Ç—É–≤–∞–Ω–Ω—è CORS - –¥–æ–∑–≤–æ–ª—è—î–º–æ –∑–∞–ø–∏—Ç–∏ –∑ –≤–∞—à–æ–≥–æ —Å–∞–π—Ç—É
CORS(app, resources={
    r"/api/*": {
        "origins": [
            "https://hlcuz.weblium.site",
            "http://hlcuz.weblium.site",
            "https://trademark-checker-rzdg.onrender.com",
            "*"  # –î–æ–∑–≤–æ–ª—è—î–º–æ –≤—Å—ñ –¥–æ–º–µ–Ω–∏ (–¥–ª—è —Ç–µ—Å—Ç—É–≤–∞–Ω–Ω—è)
        ],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "expose_headers": ["Content-Type"],
        "supports_credentials": True,
        "max_age": 3600
    }
})

# –Ü–Ω—ñ—Ü—ñ–∞–ª—ñ–∑–∞—Ü—ñ—è OpenAI –∫–ª—ñ—î–Ω—Ç–∞
try:
    api_key = os.getenv('OPENAI_API_KEY')
    if api_key:
        client = OpenAI(api_key=api_key)
    else:
        client = None
        print("Warning: OPENAI_API_KEY not set")
except Exception as e:
    print(f"Warning: OpenAI client initialization error: {e}")
    client = None

class InstructionManager:
    def __init__(self, google_doc_url):
        self.doc_url = google_doc_url
        self.cache = {}
        self.cache_expiry = None
        
    def get_instructions(self):
        if self.cache_expiry and datetime.now() < self.cache_expiry:
            return self.cache
            
        try:
            doc_id = self.extract_doc_id(self.doc_url)
            if not doc_id:
                raise Exception("–ù–µ–ø—Ä–∞–≤–∏–ª—å–Ω–∏–π URL Google Docs")
                
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            response = requests.get(export_url)
            response.raise_for_status()
            
            instructions = response.text
            
            self.cache = {
                'content': instructions,
                'updated': datetime.now()
            }
            self.cache_expiry = datetime.now() + timedelta(hours=1)
            
            return self.cache
        except Exception as e:
            print(f"–ü–æ–º–∏–ª–∫–∞ –∑–∞–≤–∞–Ω—Ç–∞–∂–µ–Ω–Ω—è —ñ–Ω—Å—Ç—Ä—É–∫—Ü—ñ–π: {e}")
            return self.cache if self.cache else {
                'content': '–í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É–π—Ç–µ –∑–∞–≥–∞–ª—å–Ω—ñ –ø—Ä–∏–Ω—Ü–∏–ø–∏ –∞–Ω–∞–ª—ñ–∑—É —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫',
                'updated': datetime.now()
            }
    
    def extract_doc_id(self, url):
        if not url:
            return None
        match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', url)
        return match.group(1) if match else None

instruction_manager = InstructionManager(os.getenv('GOOGLE_DOC_URL', ''))

# –ì–ª–æ–±–∞–ª—å–Ω–µ —Å—Ö–æ–≤–∏—â–µ –¥–ª—è —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ñ–≤ –∞–Ω–∞–ª—ñ–∑—É
analysis_storage = {}

@app.route('/')
def index():
    html_code = """
    <!DOCTYPE html>
    <html lang="uk">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>–ê–Ω–∞–ª—ñ–∑ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫</title>
        <style>
            * { margin: 0; padding: 0; box-sizing: border-box; }
            body { font-family: Arial, sans-serif; background: #f5f5f5; }
            .tm-analyzer { max-width: 1200px; margin: 0 auto; padding: 20px; }
            h1 { color: #333; margin-bottom: 30px; }
            .form-section { background: white; padding: 25px; margin: 20px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
            .form-group { margin-bottom: 15px; }
            .form-group label { display: block; margin-bottom: 5px; font-weight: bold; color: #555; }
            .form-group input, .form-group textarea { width: 100%; padding: 10px; border: 1px solid #ddd; border-radius: 4px; font-size: 14px; }
            .existing-tm { border: 2px solid #007bff; margin: 15px 0; padding: 20px; border-radius: 5px; background: #f0f8ff; }
            .btn { padding: 12px 24px; border: none; border-radius: 4px; cursor: pointer; font-size: 16px; margin: 5px; transition: 0.3s; }
            .btn:hover { opacity: 0.9; }
            .btn-primary { background: #007bff; color: white; }
            .btn-secondary { background: #6c757d; color: white; }
            .btn-success { background: #28a745; color: white; }
            .loading { text-align: center; padding: 40px; }
            .spinner { border: 4px solid #f3f3f3; border-top: 4px solid #3498db; border-radius: 50%; width: 50px; height: 50px; animation: spin 1s linear infinite; margin: 0 auto; }
            @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
            .results { margin-top: 30px; }
            .result-card { background: white; border: 1px solid #ddd; margin: 15px 0; padding: 20px; border-radius: 8px; }
            .risk-high { border-left: 5px solid #dc3545; }
            .risk-medium { border-left: 5px solid #ffc107; }
            .risk-low { border-left: 5px solid #28a745; }
            .percentage { font-size: 32px; font-weight: bold; color: #007bff; }
            .final-conclusion { background: #e8f5e8; border: 2px solid #4caf50; padding: 25px; border-radius: 8px; margin: 20px 0; }
            .success-chance { font-size: 28px; font-weight: bold; text-align: center; margin: 20px 0; }
            .tm-image { max-width: 200px; max-height: 200px; border: 1px solid #ddd; border-radius: 4px; margin: 10px 0; }
            .tm-images-container { display: flex; gap: 20px; flex-wrap: wrap; align-items: center; margin: 15px 0; }
            .image-preview { text-align: center; }
            .image-preview img { max-width: 150px; max-height: 150px; border: 2px solid #007bff; border-radius: 4px; }
            .image-preview p { margin-top: 5px; font-size: 12px; color: #666; }
            .export-buttons { text-align: center; margin: 20px 0; }
        </style>
    </head>
    <body>
        <div class="tm-analyzer">
            <h1>üîç –ê–Ω–∞–ª—ñ–∑–∞—Ç–æ—Ä —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫</h1>
            
            <form id="tmAnalyzerForm">
                <div class="form-section">
                    <h2>üìù –ë–∞–∂–∞–Ω–∞ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∞ –º–∞—Ä–∫–∞</h2>
                    <div class="form-group">
                        <label for="desired-name">–ù–∞–∑–≤–∞ *</label>
                        <input type="text" id="desired-name" required>
                    </div>
                    <div class="form-group">
                        <label for="desired-description">–û–ø–∏—Å</label>
                        <textarea id="desired-description" rows="3"></textarea>
                    </div>
                    <div class="form-group">
                        <label for="desired-classes">–ö–ª–∞—Å–∏ –ú–ö–¢–ü</label>
                        <input type="text" id="desired-classes" placeholder="25, 35, 42">
                    </div>
                    <div class="form-group">
                        <label for="desired-image">–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏</label>
                        <input type="file" id="desired-image" accept="image/*" onchange="previewImage(this, 'desired-preview')">
                        <div id="desired-preview" class="image-preview" style="display:none; margin-top:10px;"></div>
                        <p style="font-size: 12px; color: #28a745; margin-top: 5px;">
                            ‚úÖ –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –±—É–¥—É—Ç—å –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –ø—Ä–æ–∞–Ω–∞–ª—ñ–∑–æ–≤–∞–Ω—ñ –∑–∞ –¥–æ–ø–æ–º–æ–≥–æ—é GPT-4 Vision
                        </p>
                    </div>
                </div>
                
                <div class="form-section">
                    <h2>üìã –ó–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω—ñ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω—ñ –º–∞—Ä–∫–∏</h2>
                    <div id="existing-trademarks"></div>
                    <button type="button" class="btn btn-secondary" onclick="addExistingTM()">‚ûï –î–æ–¥–∞—Ç–∏ –¢–ú</button>
                </div>
                
                <div style="text-align: center;">
                    <button type="submit" class="btn btn-primary">üîç –ü—Ä–æ–≤–µ—Å—Ç–∏ –∞–Ω–∞–ª—ñ–∑</button>
                </div>
            </form>
            
            <div id="results" class="results" style="display: none;">
                <div id="loading" class="loading">
                    <div class="spinner"></div>
                    <p>–ê–Ω–∞–ª—ñ–∑—É—î–º–æ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω—ñ –º–∞—Ä–∫–∏...</p>
                </div>
                <div id="analysis-results" style="display: none;"></div>
            </div>
        </div>

        <script>
            let existingTMCount = 0;
            let analysisId = null;
            
            function previewImage(input, previewId) {
                const preview = document.getElementById(previewId);
                if (input.files && input.files[0]) {
                    const reader = new FileReader();
                    reader.onload = function(e) {
                        preview.innerHTML = `<img src="${e.target.result}" alt="–ü–æ–ø–µ—Ä–µ–¥–Ω—ñ–π –ø–µ—Ä–µ–≥–ª—è–¥"><p>–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –∑–∞–≤–∞–Ω—Ç–∞–∂–µ–Ω–æ</p>`;
                        preview.style.display = 'block';
                    }
                    reader.readAsDataURL(input.files[0]);
                } else {
                    preview.style.display = 'none';
                }
            }
            
            function addExistingTM() {
                existingTMCount++;
                const container = document.getElementById('existing-trademarks');
                const tmDiv = document.createElement('div');
                tmDiv.className = 'existing-tm';
                tmDiv.innerHTML = `
                    <h3>–¢–ú #${existingTMCount}</h3>
                    <div class="form-group">
                        <label>–ù–æ–º–µ—Ä –∑–∞—è–≤–∫–∏</label>
                        <input type="text" name="existing-${existingTMCount}-number">
                    </div>
                    <div class="form-group">
                        <label>–í–ª–∞—Å–Ω–∏–∫</label>
                        <input type="text" name="existing-${existingTMCount}-owner">
                    </div>
                    <div class="form-group">
                        <label>–ù–∞–∑–≤–∞ *</label>
                        <input type="text" name="existing-${existingTMCount}-name" required>
                    </div>
                    <div class="form-group">
                        <label>–ö–ª–∞—Å–∏ –ú–ö–¢–ü</label>
                        <input type="text" name="existing-${existingTMCount}-classes">
                    </div>
                    <div class="form-group">
                        <label>–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è</label>
                        <input type="file" name="existing-${existingTMCount}-image" accept="image/*" onchange="previewImage(this, 'existing-${existingTMCount}-preview')">
                        <div id="existing-${existingTMCount}-preview" class="image-preview" style="display:none; margin-top:10px;"></div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="removeTM(this)">‚ùå –í–∏–¥–∞–ª–∏—Ç–∏</button>
                `;
                container.appendChild(tmDiv);
            }
            
            function removeTM(button) { button.parentElement.remove(); }
            
            addExistingTM();
            
            async function fileToBase64(file) {
                return new Promise((resolve, reject) => {
                    const reader = new FileReader();
                    reader.onload = () => resolve(reader.result);
                    reader.onerror = reject;
                    reader.readAsDataURL(file);
                });
            }
            
            document.getElementById('tmAnalyzerForm').addEventListener('submit', async function(e) {
                e.preventDefault();
                document.getElementById('results').style.display = 'block';
                document.getElementById('loading').style.display = 'block';
                document.getElementById('analysis-results').style.display = 'none';
                
                const formData = new FormData(e.target);
                
                let desiredImage = null;
                const desiredImageFile = document.getElementById('desired-image').files[0];
                if (desiredImageFile) {
                    desiredImage = await fileToBase64(desiredImageFile);
                }
                
                const data = {
                    desired_trademark: {
                        name: document.getElementById('desired-name').value,
                        description: document.getElementById('desired-description').value,
                        classes: document.getElementById('desired-classes').value,
                        image: desiredImage
                    },
                    existing_trademarks: []
                };
                
                for (let i = 1; i <= existingTMCount; i++) {
                    const name = formData.get(`existing-${i}-name`);
                    if (name) {
                        let existingImage = null;
                        const existingImageInput = document.querySelector(`input[name="existing-${i}-image"]`);
                        if (existingImageInput && existingImageInput.files[0]) {
                            existingImage = await fileToBase64(existingImageInput.files[0]);
                        }
                        
                        data.existing_trademarks.push({
                            application_number: formData.get(`existing-${i}-number`) || '',
                            owner: formData.get(`existing-${i}-owner`) || '',
                            name: name,
                            classes: formData.get(`existing-${i}-classes`) || '',
                            image: existingImage
                        });
                    }
                }
                
                try {
                    const response = await fetch('/api/analyze', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(data)
                    });
                    
                    if (!response.ok) throw new Error(`HTTP ${response.status}`);
                    
                    const results = await response.json();
                    analysisId = results.analysis_id;
                    
                    document.getElementById('loading').style.display = 'none';
                    displayResults(results);
                } catch (error) {
                    document.getElementById('loading').innerHTML = `<p style="color: red;">–ü–æ–º–∏–ª–∫–∞: ${error.message}</p>`;
                }
            });
            
            function displayResults(results) {
                const container = document.getElementById('analysis-results');
                let html = '<h2>üìä –†–µ–∑—É–ª—å—Ç–∞—Ç–∏ –∞–Ω–∞–ª—ñ–∑—É</h2>';
                
                // –ó–±–µ—Ä—ñ–≥–∞—î–º–æ analysisId –≥–ª–æ–±–∞–ª—å–Ω–æ
                if (results.analysis_id) {
                    window.currentAnalysisId = results.analysis_id;
                }
                
                html += `
                    <div class="result-card" style="background: #f0f8ff; border-left: 5px solid #007bff;">
                        <h3>üéØ –ë–∞–∂–∞–Ω–∞ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∞ –º–∞—Ä–∫–∞</h3>
                        <div class="tm-images-container">
                            <div>
                                <p><strong>–ù–∞–∑–≤–∞:</strong> ${results.desired_trademark.name}</p>
                                <p><strong>–û–ø–∏—Å:</strong> ${results.desired_trademark.description || '–ù–µ –≤–∫–∞–∑–∞–Ω–æ'}</p>
                                <p><strong>–ö–ª–∞—Å–∏ –ú–ö–¢–ü:</strong> ${results.desired_trademark.classes || '–ù–µ –≤–∫–∞–∑–∞–Ω–æ'}</p>
                            </div>
                            ${results.desired_trademark.image ? `
                                <div class="image-preview">
                                    <img src="${results.desired_trademark.image}" class="tm-image" alt="–ë–∞–∂–∞–Ω–∞ –¢–ú">
                                </div>
                            ` : ''}
                        </div>
                    </div>
                `;
                
                results.results.forEach((result, index) => {
                    const riskClass = result.overall_risk > 60 ? 'risk-high' : result.overall_risk > 30 ? 'risk-medium' : 'risk-low';
                    html += `
                        <div class="result-card ${riskClass}">
                            <h3>üìÑ –ü–æ—Ä—ñ–≤–Ω—è–Ω–Ω—è –∑ –¢–ú ‚Ññ${result.trademark_info.application_number || (index + 1)}</h3>
                            
                            <div class="tm-images-container">
                                <div style="flex: 1;">
                                    <p><strong>–í–ª–∞—Å–Ω–∏–∫:</strong> ${result.trademark_info.owner}</p>
                                    <p><strong>–ù–∞–∑–≤–∞:</strong> ${result.trademark_info.name}</p>
                                    <p><strong>–ö–ª–∞—Å–∏ –ú–ö–¢–ü:</strong> ${result.trademark_info.classes}</p>
                                    <div class="percentage" style="margin-top: 15px;">${result.overall_risk}%</div>
                                    <p>–†–∏–∑–∏–∫ –∑–º—ñ—à—É–≤–∞–Ω–Ω—è: <strong>${result.confusion_likelihood}</strong></p>
                                </div>
                                ${result.trademark_info.image ? `
                                    <div class="image-preview">
                                        <img src="${result.trademark_info.image}" class="tm-image" alt="–ó–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–∞ –¢–ú">
                                        <p>–ó–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–∞ –¢–ú</p>
                                    </div>
                                ` : ''}
                            </div>
                            
                            ${result.similarity_analysis && result.similarity_analysis.phonetic ? `
                                <div style="margin: 10px 0; padding: 10px; background: #f8f9fa; border-radius: 5px;">
                                    <strong>üîä –§–æ–Ω–µ—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å:</strong> ${result.similarity_analysis.phonetic.percentage}%
                                    <p>${result.similarity_analysis.phonetic.details}</p>
                                </div>
                            ` : ''}
                            
                            ${result.similarity_analysis && result.similarity_analysis.semantic ? `
                                <div style="margin: 10px 0; padding: 10px; background: #f8f9fa; border-radius: 5px;">
                                    <strong>üí≠ –°–µ–º–∞–Ω—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å:</strong> ${result.similarity_analysis.semantic.percentage}%
                                    <p>${result.similarity_analysis.semantic.details}</p>
                                </div>
                            ` : ''}
                            
                            ${result.recommendations && result.recommendations.length > 0 ? `
                                <div style="margin: 10px 0; padding: 10px; background: #fff3e0; border-radius: 5px;">
                                    <strong>üí° –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—ó:</strong>
                                    <ul style="margin-left: 20px; margin-top: 5px;">
                                        ${result.recommendations.map(rec => `<li>${rec}</li>`).join('')}
                                    </ul>
                                </div>
                            ` : ''}
                        </div>
                    `;
                });
                
                const chanceColor = results.overall_chance > 70 ? '#4caf50' : results.overall_chance > 40 ? '#ff9800' : '#f44336';
                html += `
                    <div class="final-conclusion">
                        <h2>üìã –ó–∞–≥–∞–ª—å–Ω–∏–π –≤–∏—Å–Ω–æ–≤–æ–∫</h2>
                        <div class="success-chance" style="color: ${chanceColor}">
                            ‚úÖ –®–∞–Ω—Å —É—Å–ø—ñ—à–Ω–æ—ó —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó: ${results.overall_chance}%
                        </div>
                        <p style="text-align: center; margin-top: 10px;">
                            <small>–î–∞—Ç–∞ –∞–Ω–∞–ª—ñ–∑—É: ${new Date(results.analysis_date).toLocaleString('uk-UA')}</small>
                        </p>
                    </div>
                `;
                
                // –û–ë–û–í'–Ø–ó–ö–û–í–û –¥–æ–¥–∞—î–º–æ –∫–Ω–æ–ø–∫–∏ –µ–∫—Å–ø–æ—Ä—Ç—É
                html += `
                    <div class="export-buttons" style="margin: 30px 0; padding: 20px; background: white; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                        <h3 style="text-align: center; margin-bottom: 20px;">üì• –ó–∞–≤–∞–Ω—Ç–∞–∂–∏—Ç–∏ –∑–≤—ñ—Ç</h3>
                        <div style="display: flex; justify-content: center; gap: 15px; flex-wrap: wrap;">
                            <button class="btn btn-success" onclick="exportReport('docx')" style="font-size: 16px; padding: 15px 30px;">
                                üìÑ –ó–∞–≤–∞–Ω—Ç–∞–∂–∏—Ç–∏ DOCX
                            </button>
                            <button class="btn btn-success" onclick="exportReport('pdf')" style="font-size: 16px; padding: 15px 30px;">
                                üìë –ó–∞–≤–∞–Ω—Ç–∞–∂–∏—Ç–∏ PDF
                            </button>
                        </div>
                        <p style="text-align: center; margin-top: 15px; font-size: 14px; color: #666;">
                            –ó–≤—ñ—Ç –º—ñ—Å—Ç–∏—Ç—å –≤—Å—ñ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∏ –∞–Ω–∞–ª—ñ–∑—É —Ç–∞ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫
                        </p>
                    </div>
                `;
                
                container.innerHTML = html;
                container.style.display = 'block';
                
                // –õ–æ–≥—É–≤–∞–Ω–Ω—è –¥–ª—è –¥—ñ–∞–≥–Ω–æ—Å—Ç–∏–∫–∏
                console.log('‚úÖ –†–µ–∑—É–ª—å—Ç–∞—Ç–∏ –≤—ñ–¥–æ–±—Ä–∞–∂–µ–Ω–æ');
                console.log('üìä Analysis ID:', window.currentAnalysisId);
            }
            
            function exportReport(format) {
                const id = window.currentAnalysisId || analysisId;
                
                if (!id) {
                    alert('–ü–æ–º–∏–ª–∫–∞: ID –∞–Ω–∞–ª—ñ–∑—É –Ω–µ –∑–Ω–∞–π–¥–µ–Ω–æ. –°–ø—Ä–æ–±—É–π—Ç–µ –ø—Ä–æ–≤–µ—Å—Ç–∏ –∞–Ω–∞–ª—ñ–∑ —â–µ —Ä–∞–∑.');
                    console.error('analysisId –Ω–µ –≤—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ');
                    return;
                }
                
                console.log(`–ï–∫—Å–ø–æ—Ä—Ç —É ${format}, ID: ${id}`);
                window.location.href = `/api/export/${format}/${id}`;
            }
        </script>
    </body>
    </html>
    """
    return render_template_string(html_code)

@app.route('/api/analyze', methods=['POST', 'OPTIONS'])
def analyze_trademarks():
    # –û–±—Ä–æ–±–∫–∞ preflight OPTIONS –∑–∞–ø–∏—Ç—É
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        return response, 200
        
    try:
        data = request.json
        instructions = instruction_manager.get_instructions()
        
        results = []
        for existing_tm in data['existing_trademarks']:
            analysis = analyze_single_pair(
                desired_tm=data['desired_trademark'],
                existing_tm=existing_tm,
                instructions=instructions['content']
            )
            results.append(analysis)
        
        overall_chance = calculate_registration_chance(results)
        
        analysis_id = datetime.now().strftime('%Y%m%d%H%M%S')
        
        analysis_storage[analysis_id] = {
            'desired_trademark': data['desired_trademark'],
            'results': results,
            'overall_chance': overall_chance,
            'analysis_date': datetime.now().isoformat()
        }
        
        return jsonify({
            'analysis_id': analysis_id,
            'desired_trademark': data['desired_trademark'],
            'results': results,
            'overall_chance': overall_chance,
            'analysis_date': datetime.now().isoformat()
        })
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<format>/<analysis_id>')
def export_report(format, analysis_id):
    if analysis_id not in analysis_storage:
        return jsonify({'error': '–ê–Ω–∞–ª—ñ–∑ –Ω–µ –∑–Ω–∞–π–¥–µ–Ω–æ'}), 404
    
    analysis_data = analysis_storage[analysis_id]
    
    if format == 'docx':
        return export_docx(analysis_data, analysis_id)
    elif format == 'pdf':
        return export_pdf(analysis_data, analysis_id)
    else:
        return jsonify({'error': '–ù–µ–≤—ñ–¥–æ–º–∏–π —Ñ–æ—Ä–º–∞—Ç'}), 400

def export_docx(analysis_data, analysis_id):
    doc = Document()
    
    title = doc.add_heading('–ó–í–Ü–¢ –ü–†–û –ê–ù–ê–õ–Ü–ó –¢–û–†–ì–û–í–ï–õ–¨–ù–û–á –ú–ê–†–ö–ò', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"–î–∞—Ç–∞ –∞–Ω–∞–ª—ñ–∑—É: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()
    
    doc.add_heading('1. –ë–ê–ñ–ê–ù–ê –î–õ–Ø –†–ï–Ñ–°–¢–†–ê–¶–Ü–á –¢–û–†–ì–û–í–ï–õ–¨–ù–ê –ú–ê–†–ö–ê', 1)
    desired = analysis_data['desired_trademark']
    
    doc.add_paragraph(f"–ù–∞–∑–≤–∞: {desired['name']}")
    if desired.get('description'):
        doc.add_paragraph(f"–û–ø–∏—Å: {desired['description']}")
    if desired.get('classes'):
        doc.add_paragraph(f"–ö–ª–∞—Å–∏ –ú–ö–¢–ü: {desired['classes']}")
    
    if desired.get('image'):
        try:
            image_data = base64.b64decode(desired['image'].split(',')[1])
            image_stream = io.BytesIO(image_data)
            doc.add_picture(image_stream, width=Inches(2))
        except:
            doc.add_paragraph("–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –Ω–µ –≤–¥–∞–ª–æ—Å—è –¥–æ–¥–∞—Ç–∏")
    
    doc.add_page_break()
    
    doc.add_heading('2. –†–ï–ó–£–õ–¨–¢–ê–¢–ò –ü–û–†–Ü–í–ù–Ø–ù–ù–Ø –ó –ó–ê–†–ï–Ñ–°–¢–†–û–í–ê–ù–ò–ú–ò –¢–ú', 1)
    
    for idx, result in enumerate(analysis_data['results'], 1):
        tm_info = result['trademark_info']
        
        doc.add_heading(f'2.{idx}. –¢–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∞ –º–∞—Ä–∫–∞ ‚Ññ{tm_info.get("application_number", idx)}', 2)
        
        doc.add_paragraph(f"–í–ª–∞—Å–Ω–∏–∫: {tm_info['owner']}")
        doc.add_paragraph(f"–ù–∞–∑–≤–∞: {tm_info['name']}")
        doc.add_paragraph(f"–ö–ª–∞—Å–∏ –ú–ö–¢–ü: {tm_info['classes']}")
        
        if tm_info.get('image'):
            try:
                image_data = base64.b64decode(tm_info['image'].split(',')[1])
                image_stream = io.BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(2))
            except:
                doc.add_paragraph("–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –Ω–µ –≤–¥–∞–ª–æ—Å—è –¥–æ–¥–∞—Ç–∏")
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run(f"–†–ò–ó–ò–ö –ó–ú–Ü–®–£–í–ê–ù–ù–Ø: {result['overall_risk']}%").bold = True
        p.add_run(f" ({result['confusion_likelihood']})")
        
        if result.get('similarity_analysis'):
            doc.add_paragraph()
            doc.add_paragraph("–î–µ—Ç–∞–ª—å–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ —Å—Ö–æ–∂–æ—Å—Ç—ñ:")
            
            if result['similarity_analysis'].get('phonetic'):
                doc.add_paragraph(
                    f"‚Ä¢ –§–æ–Ω–µ—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {result['similarity_analysis']['phonetic']['percentage']}% - "
                    f"{result['similarity_analysis']['phonetic']['details']}",
                    style='List Bullet'
                )
            
            if result['similarity_analysis'].get('semantic'):
                doc.add_paragraph(
                    f"‚Ä¢ –°–µ–º–∞–Ω—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {result['similarity_analysis']['semantic']['percentage']}% - "
                    f"{result['similarity_analysis']['semantic']['details']}",
                    style='List Bullet'
                )
        
        if result.get('recommendations'):
            doc.add_paragraph()
            doc.add_paragraph("–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—ó:")
            for rec in result['recommendations']:
                doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    doc.add_page_break()
    doc.add_heading('3. –ó–ê–ì–ê–õ–¨–ù–ò–ô –í–ò–°–ù–û–í–û–ö', 1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        f"–®–∞–Ω—Å —É—Å–ø—ñ—à–Ω–æ—ó —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏ '{desired['name']}': "
    )
    chance_run = conclusion.add_run(f"{analysis_data['overall_chance']}%")
    chance_run.bold = True
    chance_run.font.size = Pt(16)
    
    if analysis_data['overall_chance'] > 70:
        chance_run.font.color.rgb = RGBColor(0, 128, 0)
        doc.add_paragraph("–í–∏—Å–æ–∫–∞ –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —É—Å–ø—ñ—à–Ω–æ—ó —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó.")
    elif analysis_data['overall_chance'] > 40:
        chance_run.font.color.rgb = RGBColor(255, 165, 0)
        doc.add_paragraph("–°–µ—Ä–µ–¥–Ω—è –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó. –†–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –¥–µ—Ç–∞–ª—å–Ω—ñ—à–µ –≤–∏–≤—á–∏—Ç–∏ –∫–æ–Ω—Ñ–ª—ñ–∫—Ç–Ω—ñ –¢–ú.")
    else:
        chance_run.font.color.rgb = RGBColor(255, 0, 0)
        doc.add_paragraph("–ù–∏–∑—å–∫–∞ –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó. –†–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –≤–Ω–µ—Å—Ç–∏ –∑–º—ñ–Ω–∏ –¥–æ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏.")
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return send_file(
        doc_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'–ê–Ω–∞–ª—ñ–∑_–¢–ú_{analysis_id}.docx'
    )

def export_pdf(analysis_data, analysis_id):
    """–ï–∫—Å–ø–æ—Ä—Ç —É PDF –±–µ–∑ –∫–∏—Ä–∏–ª–∏—Ü—ñ (—Ç—Ä–∞–Ω—Å–ª—ñ—Ç–µ—Ä–∞—Ü—ñ—è)"""
    buffer = io.BytesIO()
    
    # –§—É–Ω–∫—Ü—ñ—è —Ç—Ä–∞–Ω—Å–ª—ñ—Ç–µ—Ä–∞—Ü—ñ—ó
    def translit(text):
        """–¢—Ä–∞–Ω—Å–ª—ñ—Ç–µ—Ä–∞—Ü—ñ—è —É–∫—Ä–∞—ó–Ω—Å—å–∫–æ–≥–æ —Ç–µ–∫—Å—Ç—É"""
        ukr_to_lat = {
            '–ê':'A', '–ë':'B', '–í':'V', '–ì':'H', '“ê':'G', '–î':'D', '–ï':'E', '–Ñ':'Ye', 
            '–ñ':'Zh', '–ó':'Z', '–ò':'Y', '–Ü':'I', '–á':'Yi', '–ô':'Y', '–ö':'K', '–õ':'L', 
            '–ú':'M', '–ù':'N', '–û':'O', '–ü':'P', '–†':'R', '–°':'S', '–¢':'T', '–£':'U', 
            '–§':'F', '–•':'Kh', '–¶':'Ts', '–ß':'Ch', '–®':'Sh', '–©':'Shch', '–¨':'', 
            '–Æ':'Yu', '–Ø':'Ya',
            '–∞':'a', '–±':'b', '–≤':'v', '–≥':'h', '“ë':'g', '–¥':'d', '–µ':'e', '—î':'ye',
            '–∂':'zh', '–∑':'z', '–∏':'y', '—ñ':'i', '—ó':'yi', '–π':'y', '–∫':'k', '–ª':'l',
            '–º':'m', '–Ω':'n', '–æ':'o', '–ø':'p', '—Ä':'r', '—Å':'s', '—Ç':'t', '—É':'u',
            '—Ñ':'f', '—Ö':'kh', '—Ü':'ts', '—á':'ch', '—à':'sh', '—â':'shch', '—å':'',
            '—é':'yu', '—è':'ya', '‚Ç¥':'UAH', '‚Ññ':'#'
        }
        result = ''
        for char in str(text):
            result += ukr_to_lat.get(char, char)
        return result
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    story = []
    styles = getSampleStyleSheet()
    
    # –°—Ç–∏–ª—ñ
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=22,
        textColor=colors.HexColor('#1a237e'),
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#0d47a1'),
        spaceAfter=12,
        spaceBefore=15
    )
    
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    story.append(Paragraph(translit('–ó–í–Ü–¢ –ü–†–û –ê–ù–ê–õ–Ü–ó –¢–û–†–ì–û–í–ï–õ–¨–ù–û–á –ú–ê–†–ö–ò'), title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Data: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph('='*60, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # –ë–∞–∂–∞–Ω–∞ –¢–ú
    story.append(Paragraph(translit('1. –ë–ê–ñ–ê–ù–ê –¢–û–†–ì–û–í–ï–õ–¨–ù–ê –ú–ê–†–ö–ê'), heading_style))
    desired = analysis_data['desired_trademark']
    
    story.append(Paragraph(f"<b>Nazva:</b> {translit(desired['name'])}", styles['Normal']))
    if desired.get('description'):
        story.append(Paragraph(f"<b>Opys:</b> {translit(desired['description'])}", styles['Normal']))
    if desired.get('classes'):
        story.append(Paragraph(f"<b>Klasy MKTP:</b> {desired['classes']}", styles['Normal']))
    
    story.append(Spacer(1, 0.2*inch))
    
    # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è
    if desired.get('image'):
        try:
            image_data = base64.b64decode(desired['image'].split(',')[1])
            image_stream = io.BytesIO(image_data)
            img = RLImage(image_stream, width=2.5*inch, height=2.5*inch)
            story.append(img)
        except:
            pass
    
    story.append(PageBreak())
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç–∏
    story.append(Paragraph(translit('2. –†–ï–ó–£–õ–¨–¢–ê–¢–ò –ü–û–†–Ü–í–ù–Ø–ù–ù–Ø'), heading_style))
    story.append(Spacer(1, 0.2*inch))
    
    for idx, result in enumerate(analysis_data['results'], 1):
        tm_info = result['trademark_info']
        
        story.append(Paragraph(f'2.{idx}. TM #{tm_info.get("application_number", idx)}', heading_style))
        story.append(Paragraph(f"<b>Vlasnyk:</b> {translit(tm_info['owner'])}", styles['Normal']))
        story.append(Paragraph(f"<b>Nazva:</b> {translit(tm_info['name'])}", styles['Normal']))
        story.append(Paragraph(f"<b>Klasy:</b> {tm_info['classes']}", styles['Normal']))
        story.append(Spacer(1, 0.15*inch))
        
        # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è
        if tm_info.get('image'):
            try:
                image_data = base64.b64decode(tm_info['image'].split(',')[1])
                image_stream = io.BytesIO(image_data)
                img = RLImage(image_stream, width=2*inch, height=2*inch)
                story.append(img)
                story.append(Spacer(1, 0.15*inch))
            except:
                pass
        
        # –†–∏–∑–∏–∫
        risk = result['overall_risk']
        story.append(Paragraph(
            f'<para backColor="{"#d32f2f" if risk > 60 else "#f57c00" if risk > 30 else "#388e3c"}" textColor="white">'
            f'<b>RYZYK: {risk}%</b> ({translit(result.get("confusion_likelihood", ""))})'
            f'</para>',
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2*inch))
        
        # –ê–Ω–∞–ª—ñ–∑
        if result.get('similarity_analysis'):
            sim = result['similarity_analysis']
            
            if sim.get('phonetic'):
                story.append(Paragraph(f'<b>Fonetychna: {sim["phonetic"]["percentage"]}%</b>', styles['Normal']))
                story.append(Paragraph(translit(sim["phonetic"]["details"]), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('graphic'):
                story.append(Paragraph(f'<b>Hrafichna: {sim["graphic"]["percentage"]}%</b>', styles['Normal']))
                story.append(Paragraph(translit(sim["graphic"]["details"]), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('semantic'):
                story.append(Paragraph(f'<b>Semantychna: {sim["semantic"]["percentage"]}%</b>', styles['Normal']))
                story.append(Paragraph(translit(sim["semantic"]["details"]), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('visual'):
                story.append(Paragraph(f'<b>Vizualna: {sim["visual"]["percentage"]}%</b>', styles['Normal']))
                story.append(Paragraph(translit(sim["visual"]["details"]), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—ó
        if result.get('recommendations'):
            story.append(Paragraph('<b>Rekomendatsii:</b>', styles['Normal']))
            for rec in result['recommendations']:
                story.append(Paragraph(f'- {translit(rec)}', styles['Normal']))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('- - - - - - - -', styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # –í–∏—Å–Ω–æ–≤–æ–∫
    story.append(PageBreak())
    story.append(Paragraph(translit('3. –í–ò–°–ù–û–í–û–ö'), heading_style))
    story.append(Spacer(1, 0.3*inch))
    
    chance = analysis_data['overall_chance']
    story.append(Paragraph(
        f'<para alignment="center" fontSize="18">'
        f'Shans reyestratsiyi TM "{translit(desired["name"])}":'
        f'</para>',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    story.append(Paragraph(
        f'<para alignment="center" fontSize="36" textColor="{"#388e3c" if chance > 70 else "#f57c00" if chance > 40 else "#d32f2f"}">'
        f'<b>{chance}%</b>'
        f'</para>',
        styles['Normal']
    ))
    
    doc.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'Analiz_TM_{analysis_id}.pdf'
    )
    buffer = io.BytesIO()
    
    # –ó–∞–≤–∞–Ω—Ç–∞–∂—É—î–º–æ —à—Ä–∏—Ñ—Ç DejaVu –¥–ª—è –∫–∏—Ä–∏–ª–∏—Ü—ñ
    try:
        # –ó–∞–≤–∞–Ω—Ç–∞–∂—É—î–º–æ DejaVu Sans –∑ CDN
        dejavu_url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
        dejavu_bold_url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans-Bold.ttf"
        
        # –ó–∞–≤–∞–Ω—Ç–∞–∂—É—î–º–æ —à—Ä–∏—Ñ—Ç–∏
        dejavu_data = urllib.request.urlopen(dejavu_url).read()
        dejavu_bold_data = urllib.request.urlopen(dejavu_bold_url).read()
        
        # –ó–±–µ—Ä—ñ–≥–∞—î–º–æ —Ç–∏–º—á–∞—Å–æ–≤–æ
        with open('/tmp/DejaVuSans.ttf', 'wb') as f:
            f.write(dejavu_data)
        with open('/tmp/DejaVuSans-Bold.ttf', 'wb') as f:
            f.write(dejavu_bold_data)
        
        # –†–µ—î—Å—Ç—Ä—É—î–º–æ —à—Ä–∏—Ñ—Ç–∏
        pdfmetrics.registerFont(TTFont('DejaVu', '/tmp/DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/tmp/DejaVuSans-Bold.ttf'))
        font_name = 'DejaVu'
        font_bold = 'DejaVu-Bold'
    except:
        print("‚ö†Ô∏è –ù–µ –≤–¥–∞–ª–æ—Å—è –∑–∞–≤–∞–Ω—Ç–∞–∂–∏—Ç–∏ DejaVu, –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î–º–æ Helvetica")
        font_name = 'Helvetica'
        font_bold = 'Helvetica-Bold'
    
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    story = []
    
    # –°—Ç–∏–ª—ñ –∑ –∫–∏—Ä–∏–ª–∏—Ü–µ—é
    title_style = ParagraphStyle(
        'Title',
        fontName=font_bold,
        fontSize=24,
        textColor=colors.HexColor('#1a237e'),
        alignment=TA_CENTER,
        spaceAfter=20,
        spaceBefore=10
    )
    
    heading1_style = ParagraphStyle(
        'Heading1',
        fontName=font_bold,
        fontSize=18,
        textColor=colors.HexColor('#0d47a1'),
        spaceAfter=15,
        spaceBefore=20
    )
    
    heading2_style = ParagraphStyle(
        'Heading2',
        fontName=font_bold,
        fontSize=14,
        textColor=colors.HexColor('#1565c0'),
        spaceAfter=10,
        spaceBefore=15
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        fontName=font_name,
        fontSize=11,
        leading=16,
        spaceAfter=8
    )
    
    bold_style = ParagraphStyle(
        'Bold',
        fontName=font_bold,
        fontSize=11,
        leading=16,
        spaceAfter=8
    )
    
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫ –∑–≤—ñ—Ç—É
    story.append(Paragraph('–ó–í–Ü–¢ –ü–†–û –ê–ù–ê–õ–Ü–ó –¢–û–†–ì–û–í–ï–õ–¨–ù–û–á –ú–ê–†–ö–ò', title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # –î–∞—Ç–∞
    date_text = f"–î–∞—Ç–∞ –∞–Ω–∞–ª—ñ–∑—É: {datetime.now().strftime('%d.%m.%Y –æ %H:%M')}"
    story.append(Paragraph(date_text, normal_style))
    story.append(Spacer(1, 0.3*inch))
    
    # –õ—ñ–Ω—ñ—è-—Ä–æ–∑–¥—ñ–ª—é–≤–∞—á
    story.append(Paragraph('<para alignment="center">‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ</para>', normal_style))
    story.append(Spacer(1, 0.3*inch))
    
    # 1. –ë–ê–ñ–ê–ù–ê –¢–ú
    story.append(Paragraph('1. –ë–ê–ñ–ê–ù–ê –î–õ–Ø –†–ï–Ñ–°–¢–†–ê–¶–Ü–á –¢–û–†–ì–û–í–ï–õ–¨–ù–ê –ú–ê–†–ö–ê', heading1_style))
    
    desired = analysis_data['desired_trademark']
    
    # –¢–∞–±–ª–∏—Ü—è –∑ —ñ–Ω—Ñ–æ—Ä–º–∞—Ü—ñ—î—é
    data_table = [
        [Paragraph('<b>–ù–∞–∑–≤–∞:</b>', bold_style), Paragraph(desired['name'], normal_style)],
        [Paragraph('<b>–û–ø–∏—Å:</b>', bold_style), Paragraph(desired.get('description') or '–ù–µ –≤–∫–∞–∑–∞–Ω–æ', normal_style)],
        [Paragraph('<b>–ö–ª–∞—Å–∏ –ú–ö–¢–ü:</b>', bold_style), Paragraph(desired.get('classes') or '–ù–µ –≤–∫–∞–∑–∞–Ω–æ', normal_style)],
    ]
    
    table = Table(data_table, colWidths=[2*inch, 4.5*inch])
    table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.2*inch))
    
    # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –±–∞–∂–∞–Ω–æ—ó –¢–ú
    if desired.get('image'):
        try:
            image_data = base64.b64decode(desired['image'].split(',')[1])
            image_stream = io.BytesIO(image_data)
            img = RLImage(image_stream, width=2.5*inch, height=2.5*inch)
            story.append(Paragraph('<para alignment="center"><b>–ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏:</b></para>', bold_style))
            story.append(Spacer(1, 0.1*inch))
            story.append(img)
        except Exception as e:
            print(f"–ü–æ–º–∏–ª–∫–∞ –¥–æ–¥–∞–≤–∞–Ω–Ω—è –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è: {e}")
    
    story.append(PageBreak())
    
    # 2. –†–ï–ó–£–õ–¨–¢–ê–¢–ò –ü–û–†–Ü–í–ù–Ø–ù–ù–Ø
    story.append(Paragraph('2. –†–ï–ó–£–õ–¨–¢–ê–¢–ò –ü–û–†–Ü–í–ù–Ø–ù–ù–Ø –ó –ó–ê–†–ï–Ñ–°–¢–†–û–í–ê–ù–ò–ú–ò –¢–ú', heading1_style))
    story.append(Spacer(1, 0.2*inch))
    
    for idx, result in enumerate(analysis_data['results'], 1):
        tm_info = result['trademark_info']
        
        # –ü—ñ–¥–∑–∞–≥–æ–ª–æ–≤–æ–∫
        story.append(Paragraph(f'2.{idx}. –¢–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∞ –º–∞—Ä–∫–∞ ‚Ññ{tm_info.get("application_number", idx)}', heading2_style))
        
        # –Ü–Ω—Ñ–æ –ø—Ä–æ –¢–ú
        tm_data = [
            [Paragraph('<b>–í–ª–∞—Å–Ω–∏–∫:</b>', bold_style), Paragraph(tm_info['owner'], normal_style)],
            [Paragraph('<b>–ù–∞–∑–≤–∞:</b>', bold_style), Paragraph(tm_info['name'], normal_style)],
            [Paragraph('<b>–ö–ª–∞—Å–∏ –ú–ö–¢–ü:</b>', bold_style), Paragraph(tm_info['classes'], normal_style)],
        ]
        
        tm_table = Table(tm_data, colWidths=[2*inch, 4.5*inch])
        tm_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(tm_table)
        story.append(Spacer(1, 0.15*inch))
        
        # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–æ—ó –¢–ú
        if tm_info.get('image'):
            try:
                image_data = base64.b64decode(tm_info['image'].split(',')[1])
                image_stream = io.BytesIO(image_data)
                img = RLImage(image_stream, width=2*inch, height=2*inch)
                story.append(img)
                story.append(Spacer(1, 0.15*inch))
            except:
                pass
        
        # –†–ò–ó–ò–ö - —É –∫–æ–ª—å–æ—Ä–æ–≤—ñ–π —Ä–∞–º—Ü—ñ
        risk = result['overall_risk']
        risk_color = '#d32f2f' if risk > 60 else '#f57c00' if risk > 30 else '#388e3c'
        
        risk_para = Paragraph(
            f'<para alignment="center" backColor="{risk_color}" textColor="white" '
            f'leftIndent="10" rightIndent="10" spaceAfter="10" spaceBefore="10">'
            f'<b>–†–ò–ó–ò–ö –ó–ú–Ü–®–£–í–ê–ù–ù–Ø: {risk}%</b> ({result.get("confusion_likelihood", "–Ω–µ–≤—ñ–¥–æ–º–æ")})'
            f'</para>',
            bold_style
        )
        story.append(risk_para)
        story.append(Spacer(1, 0.2*inch))
        
        # –î–µ—Ç–∞–ª—å–Ω–∏–π –∞–Ω–∞–ª—ñ–∑
        story.append(Paragraph('<b>–î–µ—Ç–∞–ª—å–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ —Å—Ö–æ–∂–æ—Å—Ç—ñ:</b>', bold_style))
        story.append(Spacer(1, 0.1*inch))
        
        if result.get('similarity_analysis'):
            sim = result['similarity_analysis']
            
            if sim.get('phonetic'):
                story.append(Paragraph(
                    f'üîä <b>–§–æ–Ω–µ—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {sim["phonetic"]["percentage"]}%</b>',
                    bold_style
                ))
                story.append(Paragraph(sim["phonetic"]["details"], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('graphic'):
                story.append(Paragraph(
                    f'‚úçÔ∏è <b>–ì—Ä–∞—Ñ—ñ—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {sim["graphic"]["percentage"]}%</b>',
                    bold_style
                ))
                story.append(Paragraph(sim["graphic"]["details"], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('semantic'):
                story.append(Paragraph(
                    f'üí≠ <b>–°–µ–º–∞–Ω—Ç–∏—á–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {sim["semantic"]["percentage"]}%</b>',
                    bold_style
                ))
                story.append(Paragraph(sim["semantic"]["details"], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if sim.get('visual'):
                story.append(Paragraph(
                    f'üé® <b>–í—ñ–∑—É–∞–ª—å–Ω–∞ —Å—Ö–æ–∂—ñ—Å—Ç—å: {sim["visual"]["percentage"]}%</b>',
                    bold_style
                ))
                story.append(Paragraph(sim["visual"]["details"], normal_style))
                story.append(Spacer(1, 0.1*inch))
        
        # –°–ø–æ—Ä—ñ–¥–Ω–µ–Ω—ñ—Å—Ç—å —Ç–æ–≤–∞—Ä—ñ–≤
        if result.get('goods_services_relation'):
            goods = result['goods_services_relation']
            story.append(Paragraph(
                f'üì¶ <b>–°–ø–æ—Ä—ñ–¥–Ω–µ–Ω—ñ—Å—Ç—å —Ç–æ–≤–∞—Ä—ñ–≤/–ø–æ—Å–ª—É–≥: {"–¢–ê–ö" if goods.get("are_related") else "–ù–Ü"}</b>',
                bold_style
            ))
            story.append(Paragraph(goods.get("details", ""), normal_style))
            story.append(Spacer(1, 0.15*inch))
        
        # –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—ó
        if result.get('recommendations') and len(result['recommendations']) > 0:
            story.append(Paragraph('<b>üí° –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—ó:</b>', bold_style))
            for rec in result['recommendations']:
                story.append(Paragraph(f'‚Ä¢ {rec}', normal_style))
        
        # –†–æ–∑–¥—ñ–ª—é–≤–∞—á –º—ñ–∂ –¢–ú
        if idx < len(analysis_data['results']):
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph('<para alignment="center">‚Ä¢ ‚Ä¢ ‚Ä¢</para>', normal_style))
            story.append(Spacer(1, 0.2*inch))
    
    # 3. –í–ò–°–ù–û–í–û–ö
    story.append(PageBreak())
    story.append(Paragraph('3. –ó–ê–ì–ê–õ–¨–ù–ò–ô –í–ò–°–ù–û–í–û–ö', heading1_style))
    story.append(Spacer(1, 0.3*inch))
    
    chance = analysis_data['overall_chance']
    chance_color = '#388e3c' if chance > 70 else '#f57c00' if chance > 40 else '#d32f2f'
    
    story.append(Paragraph(
        f'<para alignment="center" fontSize="20">'
        f'–®–∞–Ω—Å —É—Å–ø—ñ—à–Ω–æ—ó —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏<br/>'
        f'<b>"{desired["name"]}"</b>:'
        f'</para>',
        normal_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    story.append(Paragraph(
        f'<para alignment="center" fontSize="36" textColor="{chance_color}">'
        f'<b>{chance}%</b>'
        f'</para>',
        bold_style
    ))
    story.append(Spacer(1, 0.3*inch))
    
    # –Ü–Ω—Ç–µ—Ä–ø—Ä–µ—Ç–∞—Ü—ñ—è
    if chance > 70:
        interpretation = "‚úÖ <b>–í–∏—Å–æ–∫–∞ –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —É—Å–ø—ñ—à–Ω–æ—ó —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó.</b> –¢–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∞ –º–∞—Ä–∫–∞ –º–∞—î —Ö–æ—Ä–æ—à—ñ —à–∞–Ω—Å–∏ –±—É—Ç–∏ –∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–æ—é –±–µ–∑ –∫–æ–Ω—Ñ–ª—ñ–∫—Ç—ñ–≤."
    elif chance > 40:
        interpretation = "‚ö†Ô∏è <b>–°–µ—Ä–µ–¥–Ω—è –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó.</b> –†–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –¥–µ—Ç–∞–ª—å–Ω—ñ—à–µ –≤–∏–≤—á–∏—Ç–∏ –∫–æ–Ω—Ñ–ª—ñ–∫—Ç–Ω—ñ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω—ñ –º–∞—Ä–∫–∏ —Ç–∞, –º–æ–∂–ª–∏–≤–æ, –≤–Ω–µ—Å—Ç–∏ –Ω–µ–∑–Ω–∞—á–Ω—ñ –∑–º—ñ–Ω–∏."
    else:
        interpretation = "‚ùå <b>–ù–∏–∑—å–∫–∞ –π–º–æ–≤—ñ—Ä–Ω—ñ—Å—Ç—å —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó.</b> –í–∏—è–≤–ª–µ–Ω–æ –∑–Ω–∞—á–Ω—ñ –∫–æ–Ω—Ñ–ª—ñ–∫—Ç–∏. –ù–∞—Å—Ç—ñ–π–Ω–æ —Ä–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –≤–Ω–µ—Å—Ç–∏ —Å—É—Ç—Ç—î–≤—ñ –∑–º—ñ–Ω–∏ –¥–æ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏."
    
    story.append(Paragraph(interpretation, normal_style))
    
    # –ì–µ–Ω–µ—Ä—É—î–º–æ PDF
    doc.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'–ê–Ω–∞–ª—ñ–∑_–¢–ú_{analysis_id}.pdf'
    )
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()
    
    # –°—Ç–≤–æ—Ä—é—î–º–æ —Å—Ç–∏–ª—ñ –∑ –ø—ñ–¥—Ç—Ä–∏–º–∫–æ—é –∫–∏—Ä–∏–ª–∏—Ü—ñ
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#000000'),
        spaceAfter=30,
        alignment=1,  # CENTER
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#000000'),
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica'
    )
    
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    story.append(Paragraph('ZVIT PRO ANALIZ TORGOVELNOI MARKY', title_style))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(f"Data analizu: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style))
    story.append(Spacer(1, 0.5*inch))
    
    # –ë–∞–∂–∞–Ω–∞ –¢–ú
    story.append(Paragraph('1. BAZHANA TORGOVELNA MARKA', heading_style))
    story.append(Spacer(1, 0.2*inch))
    
    desired = analysis_data['desired_trademark']
    story.append(Paragraph(f"<b>Nazva:</b> {desired['name']}", normal_style))
    if desired.get('description'):
        story.append(Paragraph(f"<b>Opys:</b> {desired['description']}", normal_style))
    if desired.get('classes'):
        story.append(Paragraph(f"<b>Klasy MKTP:</b> {desired['classes']}", normal_style))
    
    story.append(Spacer(1, 0.2*inch))
    
    # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –±–∞–∂–∞–Ω–æ—ó –¢–ú
    if desired.get('image'):
        try:
            image_data = base64.b64decode(desired['image'].split(',')[1])
            image_stream = io.BytesIO(image_data)
            img = RLImage(image_stream, width=2*inch, height=2*inch)
            story.append(img)
        except:
            story.append(Paragraph("Zobrazhennya ne vdalosya dodaty", normal_style))
    
    story.append(PageBreak())
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç–∏
    story.append(Paragraph('2. REZULTATY PORIVNYANNYA', heading_style))
    story.append(Spacer(1, 0.3*inch))
    
    for idx, result in enumerate(analysis_data['results'], 1):
        tm_info = result['trademark_info']
        
        sub_heading = ParagraphStyle('SubHead', parent=heading_style, fontSize=14)
        story.append(Paragraph(f'2.{idx}. TM #{tm_info.get("application_number", idx)}', sub_heading))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph(f"<b>Vlasnyk:</b> {tm_info['owner']}", normal_style))
        story.append(Paragraph(f"<b>Nazva:</b> {tm_info['name']}", normal_style))
        story.append(Paragraph(f"<b>Klasy:</b> {tm_info['classes']}", normal_style))
        story.append(Spacer(1, 0.1*inch))
        
        # –ó–æ–±—Ä–∞–∂–µ–Ω–Ω—è –∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–æ—ó –¢–ú
        if tm_info.get('image'):
            try:
                image_data = base64.b64decode(tm_info['image'].split(',')[1])
                image_stream = io.BytesIO(image_data)
                img = RLImage(image_stream, width=2*inch, height=2*inch)
                story.append(img)
                story.append(Spacer(1, 0.1*inch))
            except:
                pass
        
        # –†–∏–∑–∏–∫
        story.append(Paragraph(
            f"<b>RYZYK ZMISHUVANNYA: {result['overall_risk']}%</b> ({result['confusion_likelihood']})",
            normal_style
        ))
        story.append(Spacer(1, 0.2*inch))
        
        # –ê–Ω–∞–ª—ñ–∑
        if result.get('similarity_analysis'):
            story.append(Paragraph("<b>Detalnyi analiz:</b>", normal_style))
            
            if result['similarity_analysis'].get('phonetic'):
                story.append(Paragraph(
                    f"‚Ä¢ Fonetychna: {result['similarity_analysis']['phonetic']['percentage']}%",
                    normal_style
                ))
            
            if result['similarity_analysis'].get('semantic'):
                story.append(Paragraph(
                    f"‚Ä¢ Semantychna: {result['similarity_analysis']['semantic']['percentage']}%",
                    normal_style
                ))
        
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph('_' * 80, normal_style))
        story.append(Spacer(1, 0.2*inch))
    
    # –í–∏—Å–Ω–æ–≤–æ–∫
    story.append(PageBreak())
    story.append(Paragraph('3. VYSNOVOK', heading_style))
    story.append(Spacer(1, 0.3*inch))
    
    story.append(Paragraph(
        f"Shans uspishnoyi reyestratsiyi: <b>{analysis_data['overall_chance']}%</b>",
        normal_style
    ))
    
    doc.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'Analiz_TM_{analysis_id}.pdf'
    )

def analyze_single_pair(desired_tm, existing_tm, instructions):
    """–ê–Ω–∞–ª—ñ–∑—É—î –ø–∞—Ä—É —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫, –≤–∫–ª—é—á–∞—é—á–∏ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è"""
    
    # –î—ñ–∞–≥–Ω–æ—Å—Ç–∏–∫–∞ –∑–æ–±—Ä–∞–∂–µ–Ω—å
    print(f"üîç –ê–Ω–∞–ª—ñ–∑ –ø–∞—Ä–∏: '{desired_tm.get('name')}' vs '{existing_tm.get('name')}'")
    print(f"üì∏ –ë–∞–∂–∞–Ω–∞ –¢–ú –º–∞—î –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è: {bool(desired_tm.get('image'))}")
    print(f"üì∏ –ó–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–∞ –¢–ú –º–∞—î –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è: {bool(existing_tm.get('image'))}")
    if desired_tm.get('image'):
        print(f"   –†–æ–∑–º—ñ—Ä –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –±–∞–∂–∞–Ω–æ—ó: {len(desired_tm['image'])} —Å–∏–º–≤–æ–ª—ñ–≤")
    if existing_tm.get('image'):
        print(f"   –†–æ–∑–º—ñ—Ä –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–æ—ó: {len(existing_tm['image'])} —Å–∏–º–≤–æ–ª—ñ–≤")
    
    # –î–µ—Ç–∞–ª—å–Ω–∏–π –ø—Ä–æ–º–ø—Ç –∑ —ñ–Ω—Å—Ç—Ä—É–∫—Ü—ñ—è–º–∏
    text_prompt = f"""–¢–∏ –µ–∫—Å–ø–µ—Ä—Ç –∑ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫. –ü—Ä–æ–∞–Ω–∞–ª—ñ–∑—É–π –¥–≤—ñ –º–∞—Ä–∫–∏ –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ –¥–µ—Ç–∞–ª—å–Ω–æ.

–ú–ê–†–ö–ê 1 (–±–∞–∂–∞–Ω–∞): "{desired_tm.get('name', '')}"
–ö–ª–∞—Å–∏ –ú–ö–¢–ü: {desired_tm.get('classes', '–Ω–µ –≤–∫–∞–∑–∞–Ω–æ')}
–û–ø–∏—Å: {desired_tm.get('description', '–Ω–µ –≤–∫–∞–∑–∞–Ω–æ')}

–ú–ê–†–ö–ê 2 (–∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–∞): "{existing_tm.get('name', '')}"
–í–ª–∞—Å–Ω–∏–∫: {existing_tm.get('owner', '–Ω–µ –≤–∫–∞–∑–∞–Ω–æ')}
–ö–ª–∞—Å–∏ –ú–ö–¢–ü: {existing_tm.get('classes', '–Ω–µ –≤–∫–∞–∑–∞–Ω–æ')}

=== –ö–†–ò–¢–ï–†–Ü–á –ê–ù–ê–õ–Ü–ó–£ (–î–£–ñ–ï –í–ê–ñ–õ–ò–í–û) ===
{instructions[:4000]}

=== –í–ê–ñ–õ–ò–í–û ===
–ö–æ–∂–Ω–∞ –≤—ñ–¥–ø–æ–≤—ñ–¥—å –º–∞—î –±—É—Ç–∏ –î–ï–¢–ê–õ–¨–ù–û–Æ (–º—ñ–Ω—ñ–º—É–º 3-5 —Ä–µ—á–µ–Ω—å).
–í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É–π –ö–û–ù–ö–†–ï–¢–ù–Ü –ø—Ä–∏–∫–ª–∞–¥–∏ –∑ –Ω–∞–∑–≤ –º–∞—Ä–æ–∫.
–ü–æ—è—Å–Ω–∏ –ß–û–ú–£ —Ç–∏ –ø–æ—Å—Ç–∞–≤–∏–≤ —Å–∞–º–µ —Ç–∞–∫–∏–π –≤—ñ–¥—Å–æ—Ç–æ–∫.
–û–ø–∏—à–∏ –Ø–ö —Å–ø–æ–∂–∏–≤–∞—á –º–æ–∂–µ —Å–ø—Ä–∏–π–Ω—è—Ç–∏ —Ü—ñ –º–∞—Ä–∫–∏.

–Ø–∫—â–æ —î –ó–û–ë–†–ê–ñ–ï–ù–ù–Ø - –¥–µ—Ç–∞–ª—å–Ω–æ –æ–ø–∏—à–∏:
- –ö–æ–ª—å–æ—Ä–∏ (—Ç–æ—á–Ω—ñ –Ω–∞–∑–≤–∏ –∫–æ–ª—å–æ—Ä—ñ–≤)
- –ì—Ä–∞—Ñ—ñ—á–Ω—ñ –µ–ª–µ–º–µ–Ω—Ç–∏ (—â–æ —Å–∞–º–µ –∑–æ–±—Ä–∞–∂–µ–Ω–æ)
- –°—Ç–∏–ª—å (–º—ñ–Ω—ñ–º–∞–ª—ñ–∑–º, –∫–æ—Ä–ø–æ—Ä–∞—Ç–∏–≤–Ω–∏–π, –∫—Ä–µ–∞—Ç–∏–≤–Ω–∏–π —Ç–æ—â–æ)
- –ö–æ–º–ø–æ–∑–∏—Ü—ñ—é (—è–∫ —Ä–æ–∑—Ç–∞—à–æ–≤–∞–Ω—ñ –µ–ª–µ–º–µ–Ω—Ç–∏)
- –ß–∏ –º–æ–∂–Ω–∞ —ó—Ö –ø–µ—Ä–µ–ø–ª—É—Ç–∞—Ç–∏ –≤—ñ–∑—É–∞–ª—å–Ω–æ

–í—ñ–¥–ø–æ–≤—ñ–¥—å —É JSON —Ñ–æ—Ä–º–∞—Ç—ñ:
{{"trademark_info":{{"application_number":"{existing_tm.get('application_number','')}","owner":"{existing_tm.get('owner','')}","name":"{existing_tm.get('name','')}","classes":"{existing_tm.get('classes','')}"}}, "identical_test":{{"is_identical":false,"percentage":0,"details":"–î–µ—Ç–∞–ª—å–Ω–µ –æ–±“ë—Ä—É–Ω—Ç—É–≤–∞–Ω–Ω—è (3-5 —Ä–µ—á–µ–Ω—å) —á–æ–º—É –º–∞—Ä–∫–∏ —Ç–æ—Ç–æ–∂–Ω—ñ –∞–±–æ —Ä—ñ–∑–Ω—ñ"}}, "similarity_analysis":{{"phonetic":{{"percentage":0,"details":"–î–ï–¢–ê–õ–¨–ù–ò–ô –æ–ø–∏—Å (3-5 —Ä–µ—á–µ–Ω—å): —è–∫—ñ –∑–≤—É–∫–∏ —Å–ø—ñ–≤–ø–∞–¥–∞—é—Ç—å, —è–∫—ñ –≤—ñ–¥—Ä—ñ–∑–Ω—è—é—Ç—å—Å—è, —è–∫ —Ü–µ –≤–ø–ª–∏–≤–∞—î –Ω–∞ —Å–ø—Ä–∏–π–Ω—è—Ç—Ç—è, —á–∏ –ª–µ–≥–∫–æ –ø–µ—Ä–µ–ø–ª—É—Ç–∞—Ç–∏ –ø—Ä–∏ –≤–∏–º–æ–≤—ñ"}}, "graphic":{{"percentage":0,"details":"–î–ï–¢–ê–õ–¨–ù–ò–ô –æ–ø–∏—Å (3-5 —Ä–µ—á–µ–Ω—å): —è–∫—ñ –ª—ñ—Ç–µ—Ä–∏ —Å—Ö–æ–∂—ñ, —á–∏–º –≤—ñ–¥—Ä—ñ–∑–Ω—è—î—Ç—å—Å—è –≤—ñ–∑—É–∞–ª—å–Ω–æ, —á–∏ –ª–µ–≥–∫–æ –ø–µ—Ä–µ–ø–ª—É—Ç–∞—Ç–∏ –ø—Ä–∏ —á–∏—Ç–∞–Ω–Ω—ñ, –æ—Å–æ–±–ª–∏–≤–æ—Å—Ç—ñ —à—Ä–∏—Ñ—Ç—É"}}, "semantic":{{"percentage":0,"details":"–î–ï–¢–ê–õ–¨–ù–ò–ô –æ–ø–∏—Å (3-5 —Ä–µ—á–µ–Ω—å): —â–æ –æ–∑–Ω–∞—á–∞—î –∫–æ–∂–Ω–∞ –º–∞—Ä–∫–∞, —è–∫—ñ –∞—Å–æ—Ü—ñ–∞—Ü—ñ—ó –≤–∏–∫–ª–∏–∫–∞—î, —á–∏ —î –ª–æ–≥—ñ—á–Ω–∏–π –∑–≤—è–∑–æ–∫ –º—ñ–∂ –∑–Ω–∞—á–µ–Ω–Ω—è–º–∏, —â–æ –≤—ñ–¥—á—É—î —Å–ø–æ–∂–∏–≤–∞—á"}}, "visual":{{"percentage":0,"details":"–î–ï–¢–ê–õ–¨–ù–ò–ô –æ–ø–∏—Å (5-7 —Ä–µ—á–µ–Ω—å —è–∫—â–æ —î –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è): —Ç–æ—á–Ω—ñ –∫–æ–ª—å–æ—Ä–∏, –≥—Ä–∞—Ñ—ñ—á–Ω—ñ –µ–ª–µ–º–µ–Ω—Ç–∏, –∫–æ–º–ø–æ–∑–∏—Ü—ñ—è, —Å—Ç–∏–ª—å, —á–∏ –º–æ–∂–Ω–∞ –ø–µ—Ä–µ–ø–ª—É—Ç–∞—Ç–∏ –≤—ñ–∑—É–∞–ª—å–Ω–æ. –Ø–∫—â–æ –Ω–µ–º–∞—î –∑–æ–±—Ä–∞–∂–µ–Ω—å - –Ω–∞–ø–∏—à–∏ —â–æ –∞–Ω–∞–ª—ñ–∑ –Ω–µ –ø—Ä–æ–≤–µ–¥–µ–Ω–æ"}}}}, "goods_services_relation":{{"are_related":false,"details":"–î–ï–¢–ê–õ–¨–ù–ò–ô –æ–ø–∏—Å (3-5 —Ä–µ—á–µ–Ω—å): –¥–ª—è —á–æ–≥–æ –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—é—Ç—å—Å—è —Ç–æ–≤–∞—Ä–∏/–ø–æ—Å–ª—É–≥–∏, —á–∏ –æ—Ä—ñ—î–Ω—Ç–æ–≤–∞–Ω—ñ –Ω–∞ –æ–¥–Ω—É –∞—É–¥–∏—Ç–æ—Ä—ñ—é, —á–∏ –º–æ–∂—É—Ç—å –∫–æ–Ω–∫—É—Ä—É–≤–∞—Ç–∏"}}, "overall_risk":0, "confusion_likelihood":"–Ω–∏–∑—å–∫–∞/—Å–µ—Ä–µ–¥–Ω—è/–≤–∏—Å–æ–∫–∞", "recommendations":["–ö–æ–Ω–∫—Ä–µ—Ç–Ω–∞ –¥–µ—Ç–∞–ª—å–Ω–∞ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—è 1 (2-3 —Ä–µ—á–µ–Ω–Ω—è)","–ö–æ–Ω–∫—Ä–µ—Ç–Ω–∞ –¥–µ—Ç–∞–ª—å–Ω–∞ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü—ñ—è 2 (2-3 —Ä–µ—á–µ–Ω–Ω—è)"]}}"""
    
    try:
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise Exception("OpenAI API –∫–ª—é—á –Ω–µ –Ω–∞–ª–∞—à—Ç–æ–≤–∞–Ω–∏–π")
        
        if client is None:
            temp_client = OpenAI(api_key=api_key)
        else:
            temp_client = client
        
        # –ü–µ—Ä–µ–≤—ñ—Ä—è—î–º–æ —á–∏ —î –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è
        has_desired_image = desired_tm.get('image') and len(str(desired_tm.get('image', ''))) > 100
        has_existing_image = existing_tm.get('image') and len(str(existing_tm.get('image', ''))) > 100
        
        print(f"‚úÖ –ü–µ—Ä–µ–≤—ñ—Ä–∫–∞ –∑–æ–±—Ä–∞–∂–µ–Ω—å:")
        print(f"   –ë–∞–∂–∞–Ω–∞ –¢–ú: {has_desired_image}")
        print(f"   –ó–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–∞ –¢–ú: {has_existing_image}")
        
        if has_desired_image or has_existing_image:
            print(f"üé® –í–ò–ö–û–†–ò–°–¢–û–í–£–Ñ–ú–û GPT-4o Vision –¥–ª—è –∞–Ω–∞–ª—ñ–∑—É –∑–æ–±—Ä–∞–∂–µ–Ω—å")
            # –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î–º–æ GPT-4o Vision –¥–ª—è –∞–Ω–∞–ª—ñ–∑—É –∑–æ–±—Ä–∞–∂–µ–Ω—å
            messages_content = [
                {
                    "type": "text",
                    "text": text_prompt + "\n\n–£–í–ê–ì–ê: –¢–æ–±—ñ –Ω–∞–¥–∞–Ω–æ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫. –û–ë–û–í'–Ø–ó–ö–û–í–û –ø—Ä–æ–∞–Ω–∞–ª—ñ–∑—É–π —ó—Ö –≤—ñ–∑—É–∞–ª—å–Ω—É —Å—Ö–æ–∂—ñ—Å—Ç—å –¥–µ—Ç–∞–ª—å–Ω–æ!"
                }
            ]
            
            # –î–æ–¥–∞—î–º–æ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –±–∞–∂–∞–Ω–æ—ó –¢–ú (—è–∫—â–æ —î)
            if has_desired_image:
                # –ü–µ—Ä–µ–≤—ñ—Ä—è—î–º–æ —â–æ —Ü–µ data URL
                if desired_tm['image'].startswith('data:image'):
                    messages_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": desired_tm['image']
                        }
                    })
                    messages_content.append({
                        "type": "text",
                        "text": f"‚òùÔ∏è –¶–µ –ª–æ–≥–æ—Ç–∏–ø/–∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –ë–ê–ñ–ê–ù–û–á —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏ '{desired_tm.get('name', '')}'. –û–ø–∏—à–∏ –π–æ–≥–æ –¥–µ—Ç–∞–ª—å–Ω–æ."
                    })
            
            # –î–æ–¥–∞—î–º–æ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –∑–∞—Ä–µ—î—Å—Ç—Ä–æ–≤–∞–Ω–æ—ó –¢–ú (—è–∫—â–æ —î)
            if has_existing_image:
                if existing_tm['image'].startswith('data:image'):
                    messages_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": existing_tm['image']
                        }
                    })
                    messages_content.append({
                        "type": "text",
                        "text": f"‚òùÔ∏è –¶–µ –ª–æ–≥–æ—Ç–∏–ø/–∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –ó–ê–†–ï–Ñ–°–¢–†–û–í–ê–ù–û–á —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–æ—ó –º–∞—Ä–∫–∏ '{existing_tm.get('name', '')}'. –û–ø–∏—à–∏ –π–æ–≥–æ –¥–µ—Ç–∞–ª—å–Ω–æ —Ç–∞ –ø–æ—Ä—ñ–≤–Ω—è–π –∑ –ø–æ–ø–µ—Ä–µ–¥–Ω—ñ–º."
                    })
            
            # –ó–∞–ø–∏—Ç –¥–æ GPT-4o Vision
            response = temp_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "–¢–∏ –µ–∫—Å–ø–µ—Ä—Ç –∑ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫ –∑ 20-—Ä—ñ—á–Ω–∏–º –¥–æ—Å–≤—ñ–¥–æ–º. –¢–≤–æ—ó –∞–Ω–∞–ª—ñ–∑–∏ –∑–∞–≤–∂–¥–∏ –î–ï–¢–ê–õ–¨–ù–Ü —Ç–∞ –û–ë“ê–†–£–ù–¢–û–í–ê–ù–Ü. –¢–∏ –ø–∏—à–µ—à –º—ñ–Ω—ñ–º—É–º 3-5 —Ä–µ—á–µ–Ω—å –¥–ª—è –∫–æ–∂–Ω–æ–≥–æ –∫—Ä–∏—Ç–µ—Ä—ñ—é. –í—ñ–¥–ø–æ–≤—ñ–¥–∞–π –í–ò–ö–õ–Æ–ß–ù–û –≤–∞–ª—ñ–¥–Ω–∏–º JSON."
                    },
                    {
                        "role": "user",
                        "content": messages_content
                    }
                ],
                response_format={"type": "json_object"},
                max_tokens=8000,  # –ó–±—ñ–ª—å—à–µ–Ω–æ –¥–ª—è –¥–µ—Ç–∞–ª—å–Ω–∏—Ö –≤—ñ–¥–ø–æ–≤—ñ–¥–µ–π
                temperature=0.3  # –¢—Ä–æ—Ö–∏ –±—ñ–ª—å—à–µ –∫—Ä–µ–∞—Ç–∏–≤–Ω–æ—Å—Ç—ñ
            )
        else:
            # –ó–≤–∏—á–∞–π–Ω–∏–π —Ç–µ–∫—Å—Ç–æ–≤–∏–π –∞–Ω–∞–ª—ñ–∑ –±–µ–∑ –∑–æ–±—Ä–∞–∂–µ–Ω—å
            response = temp_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "–¢–∏ –µ–∫—Å–ø–µ—Ä—Ç –∑ —Ç–æ—Ä–≥–æ–≤–µ–ª—å–Ω–∏—Ö –º–∞—Ä–æ–∫ –∑ 20-—Ä—ñ—á–Ω–∏–º –¥–æ—Å–≤—ñ–¥–æ–º. –¢–≤–æ—ó –∞–Ω–∞–ª—ñ–∑–∏ –∑–∞–≤–∂–¥–∏ –î–ï–¢–ê–õ–¨–ù–Ü —Ç–∞ –û–ë“ê–†–£–ù–¢–û–í–ê–ù–Ü. –¢–∏ –ø–∏—à–µ—à –º—ñ–Ω—ñ–º—É–º 3-5 —Ä–µ—á–µ–Ω—å –¥–ª—è –∫–æ–∂–Ω–æ–≥–æ –∫—Ä–∏—Ç–µ—Ä—ñ—é. –í—ñ–¥–ø–æ–≤—ñ–¥–∞–π –í–ò–ö–õ–Æ–ß–ù–û –≤–∞–ª—ñ–¥–Ω–∏–º JSON."
                    },
                    {
                        "role": "user",
                        "content": text_prompt
                    }
                ],
                response_format={"type": "json_object"},
                temperature=0.3,
                max_tokens=8000
            )
        
        content = response.choices[0].message.content.strip()
        
        # –û—á–∏—â–µ–Ω–Ω—è –≤—ñ–¥ markdown
        content = content.replace("```json", "").replace("```", "").strip()
        lines = content.split('\n')
        cleaned_lines = [line for line in lines if not line.strip().startswith('//')]
        content = '\n'.join(cleaned_lines)
        
        print(f"‚úÖ GPT Response —É—Å–ø—ñ—à–Ω–∞ (–ø–µ—Ä—à—ñ 500 —Å–∏–º–≤–æ–ª—ñ–≤): {content[:500]}...")
        
        result = json.loads(content)
        
        # –î–æ–¥–∞—î–º–æ –∑–æ–±—Ä–∞–∂–µ–Ω–Ω—è –¥–æ —Ä–µ–∑—É–ª—å—Ç–∞—Ç—É
        if existing_tm.get('image'):
            result['trademark_info']['image'] = existing_tm['image']
        
        # –ü–µ—Ä–µ–≤—ñ—Ä–∫–∞ –æ–±–æ–≤'—è–∑–∫–æ–≤–∏—Ö –ø–æ–ª—ñ–≤
        if "trademark_info" not in result:
            result["trademark_info"] = {
                "application_number": existing_tm.get('application_number', ''),
                "owner": existing_tm.get('owner', ''),
                "name": existing_tm.get('name', ''),
                "classes": existing_tm.get('classes', '')
            }
        
        if "similarity_analysis" not in result:
            result["similarity_analysis"] = {
                "phonetic": {"percentage": 0, "details": "–ê–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π"},
                "graphic": {"percentage": 0, "details": "–ê–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π"},
                "semantic": {"percentage": 0, "details": "–ê–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π"},
                "visual": {"percentage": 0, "details": "–ê–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π"}
            }
        
        if "overall_risk" not in result:
            result["overall_risk"] = 50
            
        if "confusion_likelihood" not in result:
            result["confusion_likelihood"] = "—Å–µ—Ä–µ–¥–Ω—è"
            
        if "recommendations" not in result or not result["recommendations"]:
            result["recommendations"] = ["–†–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –¥–µ—Ç–∞–ª—å–Ω—ñ—à–µ –ø—Ä–æ–∞–Ω–∞–ª—ñ–∑—É–≤–∞—Ç–∏ –º–æ–∂–ª–∏–≤—ñ –∫–æ–Ω—Ñ–ª—ñ–∫—Ç–∏"]
        
        # –î–æ–¥–∞—î–º–æ –º—ñ—Ç–∫—É —â–æ –∞–Ω–∞–ª—ñ–∑ –∑–æ–±—Ä–∞–∂–µ–Ω—å –≤–∏–∫–æ–Ω–∞–Ω–æ
        if (has_desired_image or has_existing_image):
            if 'similarity_analysis' in result and 'visual' in result['similarity_analysis']:
                result['similarity_analysis']['visual']['images_analyzed'] = True
            
        return result
        
    except json.JSONDecodeError as e:
        print(f"‚ùå JSON Parse Error: {e}")
        print(f"Content that failed: {content if 'content' in locals() else 'No content'}")
        return create_default_result(existing_tm, f"–ü–æ–º–∏–ª–∫–∞ –ø–∞—Ä—Å–∏–Ω–≥—É JSON: {str(e)}")
        
    except Exception as e:
        print(f"‚ùå API Error: {e}")
        import traceback
        print(f"Full traceback: {traceback.format_exc()}")
        return create_default_result(existing_tm, str(e))

def create_default_result(existing_tm, error_msg):
    result = {
        "trademark_info": {
            "application_number": existing_tm.get('application_number', ''),
            "owner": existing_tm.get('owner', ''),
            "name": existing_tm.get('name', ''),
            "classes": existing_tm.get('classes', '')
        },
        "identical_test": {
            "is_identical": False, 
            "percentage": 0, 
            "details": "–ê–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π —á–µ—Ä–µ–∑ —Ç–µ—Ö–Ω—ñ—á–Ω—É –ø–æ–º–∏–ª–∫—É"
        },
        "similarity_analysis": {
            "phonetic": {
                "percentage": 0, 
                "details": "–§–æ–Ω–µ—Ç–∏—á–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π —á–µ—Ä–µ–∑ —Ç–µ—Ö–Ω—ñ—á–Ω—É –ø–æ–º–∏–ª–∫—É"
            },
            "graphic": {
                "percentage": 0, 
                "details": "–ì—Ä–∞—Ñ—ñ—á–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π —á–µ—Ä–µ–∑ —Ç–µ—Ö–Ω—ñ—á–Ω—É –ø–æ–º–∏–ª–∫—É"
            },
            "semantic": {
                "percentage": 0, 
                "details": "–°–µ–º–∞–Ω—Ç–∏—á–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π —á–µ—Ä–µ–∑ —Ç–µ—Ö–Ω—ñ—á–Ω—É –ø–æ–º–∏–ª–∫—É"
            },
            "visual": {
                "percentage": 0, 
                "details": "–í—ñ–∑—É–∞–ª—å–Ω–∏–π –∞–Ω–∞–ª—ñ–∑ –∑–æ–±—Ä–∞–∂–µ–Ω—å –ø–æ–∫–∏ –Ω–µ –ø—ñ–¥—Ç—Ä–∏–º—É—î—Ç—å—Å—è. –î–ª—è –∞–Ω–∞–ª—ñ–∑—É –∑–æ–±—Ä–∞–∂–µ–Ω—å –∑–≤–µ—Ä–Ω—ñ—Ç—å—Å—è –¥–æ –µ–∫—Å–ø–µ—Ä—Ç–∞."
            }
        },
        "goods_services_relation": {
            "are_related": False, 
            "details": "–ê–Ω–∞–ª—ñ–∑ —Å–ø–æ—Ä—ñ–¥–Ω–µ–Ω–æ—Å—Ç—ñ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∏–π —á–µ—Ä–µ–∑ —Ç–µ—Ö–Ω—ñ—á–Ω—É –ø–æ–º–∏–ª–∫—É"
        },
        "overall_risk": 0,
        "confusion_likelihood": "–Ω–µ–≤—ñ–¥–æ–º–æ",
        "recommendations": [
            "–°—Ç–∞–ª–∞—Å—è —Ç–µ—Ö–Ω—ñ—á–Ω–∞ –ø–æ–º–∏–ª–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª—ñ–∑—ñ",
            "–†–µ–∫–æ–º–µ–Ω–¥—É—î—Ç—å—Å—è –ø–æ–≤—Ç–æ—Ä–∏—Ç–∏ —Å–ø—Ä–æ–±—É",
            f"–î–µ—Ç–∞–ª—ñ –ø–æ–º–∏–ª–∫–∏: {error_msg}"
        ]
    }
    
    if existing_tm.get('image'):
        result['trademark_info']['image'] = existing_tm['image']
    
    return result

def calculate_registration_chance(results):
    if not results:
        return 95
    max_risk = max([result.get('overall_risk', 0) for result in results])
    if max_risk > 80:
        return 10
    elif max_risk > 60:
        return 30
    elif max_risk > 40:
        return 60
    elif max_risk > 20:
        return 80
    else:
        return 95

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
